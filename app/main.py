from fastapi import FastAPI, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import logging, io

app = FastAPI()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@app.post("/modify-docx-file/")
async def modify_docx_file(file: UploadFile):
	"""
	Accepts a .docx file as an input, adds two paragraphs to it,
	removes unnecessary leading spaces and/or empty paragraphs,
	changes all of the paragraph styles to any built-in style,
	and returns the modified .docx file as binary data in the response.
	python-docx library is used to modify the .docx file. 
	Documentation: https://python-docx.readthedocs.io/en/latest/index.html
	"""
	try:
		# Validate file extension
		if not file.filename.endswith(".docx"):
			raise HTTPException(status_code=400, detail="Invalid file type. Only .docx files are allowed.")

		# Read the contents of the uploaded .docx file in binary mode
		file_contents = await file.read()

		# Open the .docx file from the contents
		document = Document(io.BytesIO(file_contents))

		# Add two new paragraphs to the document
		document.add_paragraph("Hey Burak, thanks for the fun test, I really enjoyed it!")
		document.add_paragraph("Hope to be able to help you with your Python projects in the future.")

		# Remove unnecessary leading spaces and empty paragraphs
		for paragraph in document.paragraphs:
			if not paragraph.text.strip():
				# Remove empty paragraphs
				if paragraph._element.getparent() is not None:
					paragraph._element.getparent().remove(paragraph._element)
			else:
				# Remove leading spaces
				paragraph.text = paragraph.text.lstrip()

		# Change all paragraph styles to "Normal"
		for style in document.styles:
			if style.type == WD_STYLE_TYPE.PARAGRAPH:
				style.name = "Normal"

		# Save the modified .docx file as a bytes object
		modified_file = io.BytesIO()
		document.save(modified_file)
		modified_file.seek(0)

		# Log the successful processing of the file
		logger.info(f"File {file.filename} was successfully processed.")

		# Return the modified .docx file as binary data in the response
		return StreamingResponse(iter([modified_file.getvalue()]), 
			media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
			headers={"content-disposition": f"attachment; filename={file.filename}"})

	except Exception as e:
		# Log the error
		logger.error(f"Error processing file {file.filename}: {e}")

		# Raise HTTP exception with 500 status code and error message
		raise HTTPException(status_code=500, detail="Error processing file. Please try again later.")
