import logging
from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from fastapi import HTTPException

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileModifier:
	def modify_docx(self, file_contents: bytes) -> bytes:
		try:
			# Create a Document object from the given file contents
			document = Document(BytesIO(file_contents))

			# Call the following three methods to modify the document
			self._add_paragraphs(document)
			self._remove_unnecessary_spaces(document)
			self._change_paragraph_styles(document)

			# Save the modified document to a BytesIO object and return its value
			modified_file = BytesIO()
			document.save(modified_file)
			modified_file.seek(0)

			return modified_file.getvalue()
		except BadZipFile:
			logger.error("Error processing file BadZipFile")
			raise HTTPException(status_code=400, detail="Invalid file format. Only .docx files are allowed.")

	def _add_paragraphs(self, document: Document):
		# Add two paragraphs with specific text to the given Document object
		document.add_paragraph("Hey Burak, thanks for the fun test, I really enjoyed it!")
		document.add_paragraph("Hope to be able to help you with your Python projects in the future.")
	
	def _remove_unnecessary_spaces(self, document: Document):
		# Remove any empty paragraphs or paragraphs containing only whitespace from the given Document object
		for paragraph in document.paragraphs:
			if not paragraph.text.strip() and paragraph._element.getparent() is not None:
				paragraph._element.getparent().remove(paragraph._element)
			else:
				paragraph.text = paragraph.text.lstrip()
	
	def _change_paragraph_styles(self, document: Document):
		# Change the style of all paragraphs in the given Document object to "Normal"
		for paragraph in document.paragraphs:
			paragraph.style = document.styles["Normal"]
