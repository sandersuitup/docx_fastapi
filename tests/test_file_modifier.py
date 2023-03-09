import io

from docx import Document

from modules.file_modifier import FileModifier


def test_modify_docx():
	# Create a sample .docx file with known contents
	document = Document()
	document.add_paragraph("This is a sample paragraph.")
	file = io.BytesIO()
	document.save(file)

	# Create a FileModifier instance and modify the sample .docx file
	file_modifier = FileModifier()
	modified_file_contents = file_modifier.modify_docx(file.getvalue())

	# Open the modified file as a Document object and get the contents
	modified_document = Document(io.BytesIO(modified_file_contents))
	modified_contents = "\n".join([p.text for p in modified_document.paragraphs])

	# Check that the modified file contents are as expected
	expected_contents = "This is a sample paragraph.\nHey Burak, thanks for the fun test, I really enjoyed it!\nHope to be able to help you with your Python projects in the future."
	assert modified_contents == expected_contents
