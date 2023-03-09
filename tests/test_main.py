from io import BytesIO
from docx import Document
import sys, os
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from app.main import app

client = TestClient(app)

# Define a test case for the modify_docx_file endpoint
def test_modify_docx_file():
	# Create a fake .docx file
	test_file = BytesIO()
	document = Document()
	document.add_paragraph("Existing paragraph")
	document.save(test_file)
	test_file.seek(0)

	# Call the modify_docx_file function with the fake file
	response = client.post("/modify-docx-file/", files={"file": ("test.docx", test_file)})

	# Check that the response has a 200 status code
	assert response.status_code == 200
	
	# Check that the response contains a valid .docx file
	assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
	assert "content-disposition" in response.headers
	assert response.headers["content-disposition"].startswith("attachment; filename=")
	assert os.path.splitext(response.headers["content-disposition"])[1] == ".docx"

	# Check that the response contains the expected content
	modified_file = BytesIO(response.content)
	modified_document = Document(modified_file)
	assert len(modified_document.paragraphs) == 3
	assert modified_document.paragraphs[0].text == "Existing paragraph"
	assert modified_document.paragraphs[1].text == "Hey Burak, thanks for the fun test, I really enjoyed it!"
	assert modified_document.paragraphs[2].text == "Hope to be able to help you with your Python projects in the future."